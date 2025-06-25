from docx import Document

def write_multi_page_ocr_output_to_docx(ocr_data_pages, output_path):
    """
    Combine multiple pages of OCR data into one DOCX document, including top text and tables.
    """
    doc = Document()

    for i, ocr_data in enumerate(ocr_data_pages):
        top_text = ocr_data.get("top_text", "").strip()
        tables = ocr_data.get("tables", [])

        doc.add_heading(f"Page {i + 1}", level=1)

        if top_text:
            # doc.add_paragraph(top_text)#
            #doc.add_paragraph("")  # spacing#

        if not tables:
            doc.add_paragraph("No tables detected.")
        else:
            for t_index, table in enumerate(tables):
                data = table.get("data", [])
                if data:
                    doc_table = doc.add_table(rows=len(data), cols=len(data[0]))
                    doc_table.style = "Table Grid"
                    for r, row in enumerate(data):
                        for c, cell_text in enumerate(row):
                            doc_table.cell(r, c).text = cell_text or ""
                    doc.add_paragraph("")  # spacing after each table

        if i < len(ocr_data_pages) - 1:
            doc.add_page_break()

    doc.save(output_path)
    print(f"[DOCX] Multi-page DOCX saved to: {output_path}")
